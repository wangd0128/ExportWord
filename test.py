import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import sys
import os
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import win32com.client as win32
import time
import datetime


def read_word_tables_with_win32(file_path):
    word = None
    doc = None
    try:
        # 创建 Word 应用程序实例
        word = win32.Dispatch('Word.Application')
        word.Visible = False  # 设置为不可见
        
        # 打开文档
        doc = word.Documents.Open(file_path)
        tables_content = []
 
        for table_index, table in enumerate(doc.Tables):
            print(f"正在处理文件 {file_path} 的第 {table_index + 1} 个表格")
            rows_content = {}
            max_cols = table.Columns.Count
            
            # Initialize empty lists for all rows
            for i in range(0, table.Rows.Count):
                rows_content[i] = [''] * max_cols

            offset = 0
            prev_row_idx = 0
            # 先填充所有单元格内容
            for cell in table.Range.Cells:
                row_idx = cell.RowIndex - 1

                if row_idx == 5:
                    print(cell.Range.Text)
                if row_idx > prev_row_idx:
                    offset = 0
                    prev_row_idx = row_idx
                col_idx = cell.ColumnIndex - 1
                
                # 获取单元格文本内容并清理
                cell_text = cell.Range.Text.strip()
                # 移除特殊字符
                cell_text = cell_text.replace('\r', '').replace('\x07', '')
                if row_idx < 2:
                 
                    for i in range(cell.Range.Columns.Count):
                        rows_content[row_idx][col_idx + offset + i] = cell_text
                    offset += cell.Range.Columns.Count - 1
                else:
                    rows_content[row_idx][col_idx] = cell_text
                print(cell_text + ' : ' + str(offset) + ' ' + str(row_idx) + ' ' + str(col_idx))
                
            tables_content.append(rows_content)
        
        # 使用第一行作为列名
        first_columns = tables_content[0][0]
        second_columns = tables_content[0][1]
        columns = []
        for i in range(len(first_columns)):
            columns.append(first_columns[i] + 
                         ('（' + second_columns[i] + '）' if second_columns[i] != '' else ''))

        # 将字典数据转换为列表格式
        data_rows = []
        for i in range(2, len(tables_content[0])):
            data_rows.append(tables_content[0][i])

        df = pd.DataFrame(data_rows, columns=columns)
        
        # 清理数据：删除全空行和重复行
        df = df.dropna(how='all').drop_duplicates()
        
        return df

    except Exception as e:
        print(f"读取Word文件时出错: {str(e)}")
        raise
    finally:
        # 确保正确关闭文档和Word应用程序
        try:
            if doc:
                doc.Close()
        except Exception as e:
            print(f"关闭Word文件时出错1: {str(e)}")
            pass

        try:
            if word:
                word.Quit()
        except Exception as e:
            print(f"关闭Word文件时出错2: {str(e)}")
            pass



def clean_data(df):
    """清理数据，处理异常值和格式"""
    # 替换'/'为NaN
    df = df.replace('/', np.nan)
    # 删除完全重复的行
    df = df.drop_duplicates()
    return df

def extract_operation(col):
    """从列名中提取操作名称"""
    parts = col.split('（')
    operation_name = parts[0]
    operation_detail = parts[1].rstrip('）') if len(parts) > 1 else ''
    return operation_name, operation_detail


def save_to_word(df_tuple, output_file):
    """保存数据到Word文档"""
    df = df_tuple[0]  # 获取DataFrame
    doc = Document()
    months = [4,5,6]
    # 添加分隔线
    doc.add_paragraph('-' * 50)
    
    # 创建表格
   
   
    
    try: # 构建操作字典
        operations = {}
        for index, df in enumerate(df_tuple):
            for col in df.columns[1:]:  # 跳过姓名列
                category, operation = extract_operation(col)
                if category not in operations:
                    operations[category] = {}
            
                if months[index] not in operations[category]:
                    operations[category][months[index]] = []
                operations[category][months[index]].append(operation)
            

        # 定义类别映射
        category_mapping = {
            '理论': '理论',
            '专科操作': '专科操作',
            '基础操作': '基础操作',
            '急救操作': '急救操作',
            'OSCE': 'OSCE/综合能力',
            '综合能力': 'OSCE/综合能力'
        }
 
    
        
        # 计算每个类别的操作数量
        categories = [(cat, len(ops) if len(ops) > 0 else 1) for cat, ops in operations.items()]
        total_operations = sum(count for _, count in categories)
        print(f"Total operations count: {total_operations}")
        print("Operations per category:")
        total_cols = 1
       #len(months) * total_operations
        operation_names = list(operations.keys())
        num_operations = len(operation_names)
        for idx, operation in enumerate(operation_names):
            if idx == num_operations - 1:
                break;
            operation_dict = operations[operation]
            _len = len(months)
            if operation_dict:
                _len = sum(len(operation_dict[month]) for month in months)
            total_cols = total_cols + _len

        table = doc.add_table(rows=len(df)+4, cols=total_cols)
        table.style = 'Table Grid'
        
        # 设置表格格式
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = '宋体'
                        #居中对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


         # 处理第一行（N2理论、操作、综合能力考核）
        # 合并姓名列
        merge_cells(table.cell(0, 0), table.cell(3, 0))
        table.cell(0, 0).text = '姓名'
        
        # 合并第一行的标题（从第2列到最后一列）
        merge_cells(table.cell(0, 1), table.cell(0, total_cols-1))
        table.cell(0, 1).text = '理论、操作、综合能力考核'


        print(operations)
        # 处理第二行到第四行
        current_col = 1

        set_columns = [[] for _ in range(len(months))]
        for idx, operation in enumerate(operation_names):
            if idx == num_operations - 1:
                break;
            operation_dict = operations[operation]
            _len = len(months)
            if operation_dict:
                _len = sum(len(operation_dict[month]) for month in months)

            operation_dict_index = 0
            for item, value in operation_dict.items():
                merge_cells(table.cell(2, current_col + operation_dict_index), table.cell(2, current_col + operation_dict_index + len(value) - 1))
                table.cell(2, current_col + operation_dict_index).text = f"{item}月"
                
                # 使用列表推导式和 sum() 函数来计算 _len
                for month_value_index, month_value in enumerate(value):
                    set_columns[months.index(item)].append(current_col + operation_dict_index + month_value_index)
                    table.cell(3, current_col + operation_dict_index + month_value_index).text = month_value
                operation_dict_index += len(value)
            

            merge_cells(table.cell(1, current_col), table.cell(1, current_col + _len - 1))

            print("合并单元格 ", operation, current_col, current_col + _len - 1)
            table.cell(1, current_col).text = operation
                
                
            current_col = current_col + _len
        print(set_columns)
        _set_columns = [[1,4,7,8,13,16],[2,5,9,10,14,17],[3,6,11,12,15,18]]
        # 获取所有姓名（第一列），从第二行开始
        all_names = []
        for df in df_tuple: 
            for i, row in df.iterrows():
                if row[0] not in all_names:
                    all_names.append(row[0])
       #all_names = [all_names[0]]
        # 添加数据
        i = 0
        for index, df in enumerate(df_tuple):
            for j, row in df.iterrows():
                if  row[0] in all_names:
                    for k, value in enumerate(row):
                        _row = all_names.index(row[0]) + 4
                        if k == 0:
                            cell = table.cell(_row, 0)
                            cell.text = row[0]
                        elif k >= 1 and k <= len(set_columns[i]):
                            _col = set_columns[i][k - 1]
                            print(f"row: {_row}, col: {_col}")
                            cell = table.cell(_row, _col)
                            cell.text = str(value) if pd.notna(value) else ''
                        elif k == len(set_columns[i]) + 1:
                            _col = set_columns[i][k - 2]
                            print(f"row: {_row}, col: {_col}")
                            cell = table.cell(_row, _col)
                            if cell.text == '' or cell.text == '/':
                                cell.text = str(value) if pd.notna(value) else ''
            i += 1
        
        # 添加分隔线
        doc.add_paragraph('-' * 50)
        
        # 保存文档
        doc.save(output_file)
        
    except Exception as e:
        print(f"创建表格时出错: {str(e)}")
        raise

def merge_cells(start_cell, end_cell):
    """安全地合并单元格"""
    try:
        start_cell.merge(end_cell)
    except Exception as e:
        print(f"合并单元格时出错: {str(e)}")

def main():
    # 指定输入文件列表
    input_files = [
        r'F:\Github\ExportWord\10.docx',
        r'F:\Github\ExportWord\11.docx',
        r'F:\Github\ExportWord\12.docx'
    ]
    
    # 从多个Word文件读取表格数据
    input_tables = []
    for file in input_files:
        try:
            if os.path.exists(file):
                table_data = read_word_tables_with_win32(file)
                    
                time.sleep(1)
                if table_data is not None:
                    input_tables.append(table_data)
            else:
                print(f"警告：文件 {file} 不存在")
        except Exception as e:
            print(f"处理文件 {file} 时出错: {str(e)}")
            continue
    
    if not input_tables:
        print("错误：没有找到任何有效的表格数据")
        return
    
    # 处理数据
    try:
        # 获取当前时间并格式化为字符串
        current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        # 生成带有时间戳的文件名
        output_file_name = f'output_{current_time}.docx'
        
        # 保存到Word文档
        save_to_word(input_tables, output_file_name)
        print(f"成功生成输出文件：{output_file_name}")
        
    except Exception as e:
        print(f"处理数据时出错：{str(e)}")

if __name__ == "__main__":
    main()
