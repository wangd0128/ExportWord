import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import sys
import os
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client as win32
import time
import datetime
import gradio as gr
import pythoncom


def read_word_tables_with_win32(file_path):
    word = None
    doc = None
    try:
        # 初始化COM环境
        pythoncom.CoInitialize()
        
        # 创建 Word 应用程序实例
        word = win32.Dispatch('Word.Application')
        word.Visible = False  # 设置为不可见
        
        # 打开文档
        doc = word.Documents.Open(file_path)
        tables_content = []
 
        for table_index, table in enumerate(doc.Tables):
            print(f"正在处理文件 {file_path} 的第 {table_index + 1} 个表格")
            rows_content = {}
            max_cols = table.Columns.Count
            
            # Initialize empty lists for all rows
            for i in range(0, table.Rows.Count):
                rows_content[i] = [''] * max_cols

            offset = 0
            prev_row_idx = 0
            # 先填充所有单元格内容
            for cell in table.Range.Cells:
                row_idx = cell.RowIndex - 1

                if row_idx == 5:
                    print(cell.Range.Text)
                if row_idx > prev_row_idx:
                    offset = 0
                    prev_row_idx = row_idx
                col_idx = cell.ColumnIndex - 1
                
                # 获取单元格文本内容并清理
                cell_text = cell.Range.Text.strip()
                # 移除特殊字符
                cell_text = cell_text.replace('\r', '').replace('\x07', '')
                if row_idx < 2:
                 
                    for i in range(cell.Range.Columns.Count):
                        rows_content[row_idx][col_idx + offset + i] = cell_text
                    offset += cell.Range.Columns.Count - 1
                else:
                    rows_content[row_idx][col_idx] = cell_text
                print(cell_text + ' : ' + str(offset) + ' ' + str(row_idx) + ' ' + str(col_idx))
                
            tables_content.append(rows_content)
        
        # 使用第一行作为列名
        first_columns = tables_content[0][0]
        second_columns = tables_content[0][1]
        columns = []
        for i in range(len(first_columns)):
            columns.append(first_columns[i] + 
                         ('（' + second_columns[i] + '）' if second_columns[i] != '' else ''))

        # 将字典数据转换为列表格式
        data_rows = []
        for i in range(2, len(tables_content[0])):
            data_rows.append(tables_content[0][i])

        df = pd.DataFrame(data_rows, columns=columns)
        
        # 清理数据：删除全空行和重复行
        df = df.dropna(how='all').drop_duplicates()
        
        return df

    except Exception as e:
        print(f"读取Word文件时出错: {str(e)}")
        raise
    finally:
        # 确保正确关闭文档和Word应用程序
        try:
            if doc:
                doc.Close()
        except Exception as e:
            print(f"关闭Word文件时出错1: {str(e)}")
            pass

        try:
            if word:
                word.Quit()
        except Exception as e:
            print(f"关闭Word文件时出错2: {str(e)}")
            pass
            
        # 取消初始化COM环境
        pythoncom.CoUninitialize()



def clean_data(df):
    """清理数据，处理异常值和格式"""
    # 替换'/'为NaN
    df = df.replace('/', np.nan)
    # 删除完全重复的行
    df = df.drop_duplicates()
    return df

def extract_operation(col):
    """从列名中提取操作名称"""
    parts = col.split('（')
    operation_name = parts[0]
    operation_detail = parts[1].rstrip('）') if len(parts) > 1 else ''
    return operation_name, operation_detail


def save_to_word(df_tables, output_file):
    """保存数据到Word文档"""
    df_tuple = df_tables[0]  # 获取DataFrame
    doc = Document()
    months = df_tables[1]
   
    # 构建操作字典
    operations = {}
    for index, df in enumerate(df_tuple):
        for col in df.columns[1:]:  # 跳过姓名列
            category, operation = extract_operation(col)
            if category not in operations:
                operations[category] = {}
            
            if months[index] not in operations[category]:
                operations[category][months[index]] = []
            operations[category][months[index]].append(operation)
            
    total_cols = 1
    operation_names = list(operations.keys())
    num_operations = len(operation_names)
    for idx, operation in enumerate(operation_names):
        if idx == num_operations - 1:
            break
        operation_dict = operations[operation]
        _len = len(months)
        if operation_dict:
            _len = sum(len(operation_dict[month]) for month in months)
        total_cols = total_cols + _len

    table = doc.add_table(rows=len(df_tuple[0])+4, cols=total_cols)
    table.style = 'Table Grid'
    
    # 设置表格格式 - 修改这部分以实现居中对齐
    for row in table.rows:
        for cell in row.cells:
            # 设置单元格垂直居中
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "center")
            tcPr.append(tcVAlign)
            
            # 设置段落水平居中
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'

    # 处理第一行（N2理论、操作、综合能力考核）
    # 合并姓名列
    merge_cells(table.cell(0, 0), table.cell(3, 0))
    table.cell(0, 0).text = '姓名'
    
    # 合并第一行的标题（从第2列到最后一列）
    merge_cells(table.cell(0, 1), table.cell(0, total_cols-1))
    table.cell(0, 1).text = '理论、操作、综合能力考核'


    print(operations)
    # 处理第二行到第四行
    current_col = 1

    set_columns = [[] for _ in range(len(months))]
    for idx, operation in enumerate(operation_names):
        if idx == num_operations - 1:
            break;
        operation_dict = operations[operation]
        _len = len(months)
        if operation_dict:
            _len = sum(len(operation_dict[month]) for month in months)

        operation_dict_index = 0
        for item, value in operation_dict.items():
            merge_cells(table.cell(2, current_col + operation_dict_index), table.cell(2, current_col + operation_dict_index + len(value) - 1))
            table.cell(2, current_col + operation_dict_index).text = f"{item}月"
            
            # 使用列表推导式和 sum() 函数来计算 _len
            for month_value_index, month_value in enumerate(value):
                set_columns[months.index(item)].append(current_col + operation_dict_index + month_value_index)
                table.cell(3, current_col + operation_dict_index + month_value_index).text = month_value
            operation_dict_index += len(value)
        

        merge_cells(table.cell(1, current_col), table.cell(1, current_col + _len - 1))

        print("合并单元格 ", operation, current_col, current_col + _len - 1)
        table.cell(1, current_col).text = operation
            
            
        current_col = current_col + _len
    print(set_columns)
    _set_columns = [[1,4,7,8,13,16],[2,5,9,10,14,17],[3,6,11,12,15,18]]
    # 获取所有姓名（第一列），从第二行开始
    all_names = []
    for df in df_tuple: 
        for i, row in df.iterrows():
            if row[0] not in all_names:
                all_names.append(row[0])
   #all_names = [all_names[0]]
    # 添加数据
    i = 0
    for index, df in enumerate(df_tuple):
        for j, row in df.iterrows():
            if  row[0] in all_names:
                for k, value in enumerate(row):
                    _row = all_names.index(row[0]) + 4
                    if k == 0:
                        cell = table.cell(_row, 0)
                        cell.text = row[0]
                    elif k >= 1 and k <= len(set_columns[i]):
                        _col = set_columns[i][k - 1]
                        print(f"row: {_row}, col: {_col}")
                        cell = table.cell(_row, _col)
                        cell.text = str(value) if pd.notna(value) else ''
                    elif k == len(set_columns[i]) + 1:
                        _col = set_columns[i][k - 2]
                        print(f"row: {_row}, col: {_col}")
                        cell = table.cell(_row, _col)
                        if cell.text == '' or cell.text == '/':
                            cell.text = str(value) if pd.notna(value) else ''
        i += 1
    
    # 添加分隔线
    doc.add_paragraph('-' * 50)
    
    # 保存文档
    doc.save(output_file)
    

def merge_cells(start_cell, end_cell):
    """安全地合并单元格"""
    try:
        start_cell.merge(end_cell)
    except Exception as e:
        print(f"合并单元格时出错: {str(e)}")

def process_files(files, months):
    """处理上传的文件和对应的月份配置"""
    if len(files) != len(months):
        raise ValueError("文件数量与月份配置数量不匹配")
        
    input_tables = []
    for file, month in zip(files, months):
        try:
            if os.path.exists(file.name):
                table_data = read_word_tables_with_win32(file.name)
                time.sleep(1)
                if table_data is not None:
                    input_tables.append(table_data)
            else:
                raise FileNotFoundError(f"文件 {file.name} 不存在")
        except Exception as e:
            raise Exception(f"处理文件 {file.name} 时出错: {str(e)}")
    
    if not input_tables:
        raise ValueError("没有找到任何有效的表格数据")
    
    # 使用绝对路径保存文件
    current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(os.getcwd(), "outputs")
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    output_file_name = os.path.join(output_dir, f'output_{current_time}.docx')
    save_to_word((input_tables, months), output_file_name)
    return output_file_name

def create_ui():
    """创建Gradio界面"""
    with gr.Blocks() as app:
        gr.Markdown("""
        # Word表格处理工具
        
        ## 使用说明
        1. 选择要处理的Word文件（支持1-3个文件）
        2. 为每个文件设置对应的月份
        3. 点击"处理文件"按钮开始处理
        4. 处理完成后可以下载生成的文件
        """)
        
        with gr.Row():
            with gr.Column():
                file1 = gr.File(label="选择第一个Word文件")
                month1 = gr.Number(label="第一个文件对应月份", value=4, minimum=1, maximum=12)
            
            with gr.Column():
                file2 = gr.File(label="选择第二个Word文件")
                month2 = gr.Number(label="第二个文件对应月份", value=5, minimum=1, maximum=12)
            
            with gr.Column():
                file3 = gr.File(label="选择第三个Word文件")
                month3 = gr.Number(label="第三个文件对应月份", value=6, minimum=1, maximum=12)
        
        with gr.Row():
            submit_btn = gr.Button("处理文件", variant="primary")
            output_file = gr.File(label="下载处理后的文件")
        
        status_output = gr.Textbox(label="处理状态")

        def wrapper(f1, f2, f3, m1, m2, m3):
            files = [f for f in [f1, f2, f3] if f is not None]
            months = [m for f, m in zip([f1, f2, f3], [m1, m2, m3]) if f is not None]
            if not files:
                return "请至少选择一个文件", None
            
            try:
                output_path = process_files(files, months)
                # 确认文件存在
                if os.path.exists(output_path):
                    return f"文件处理成功！", output_path
                else:
                    return f"文件生成失败", None
            except Exception as e:
                return f"处理数据时出错：{str(e)}", None

        submit_btn.click(
            fn=wrapper,
            inputs=[file1, file2, file3, month1, month2, month3],
            outputs=[status_output, output_file]
        )
    
    return app

def main():
    app = create_ui()
    # 修改launch参数，自动打开浏览器
    app.launch(inbrowser=True, share=False)

if __name__ == "__main__":
    main()
